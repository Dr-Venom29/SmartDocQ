from __future__ import annotations

import io
import os
import sys
import types
import threading
import pytest

# Ensure backend/config.py can import in test environments.
os.environ.setdefault("SERVICE_TOKEN", "test")

# ---------------------------------------------------------------------------
# Stub external deps BEFORE importing indexer / retrieval_service
# ---------------------------------------------------------------------------

class _ImportCollectionStub:
    def upsert(self, *args, **kwargs):
        pass

    def get(self, *args, **kwargs):
        return {"ids": [], "documents": [], "metadatas": []}

    def delete(self, *args, **kwargs):
        pass

    def query(self, *args, **kwargs):
        return {"documents": [[]], "distances": [[]], "metadatas": [[]]}


fake_db_chroma = types.ModuleType("db.chroma")
fake_db_chroma.collection = _ImportCollectionStub()
sys.modules["db.chroma"] = fake_db_chroma


# No longer stubbing services.embedding_service globally as it is fully mockable.


# bm25_service is imported by retrieval_service; stub the search function only
# so the module loads without rank_bm25 needing a live corpus.
try:
    import services.bm25_service as _bm25_mod
except Exception:
    _bm25_mod = types.ModuleType("services.bm25_service")
    _bm25_mod.bm25_search = lambda *_a, **_k: []
    sys.modules["services.bm25_service"] = _bm25_mod


from indexing import indexer
import indexing.pipeline as pipeline
from services import retrieval_service
from services import bm25_service
import indexing.chunking as chunking_mod

from utils.extraction import extract_docx
from utils.table_extraction import (
    extract_tables_from_csv_bytes,
    extract_tables_from_docx_bytes,
    extract_tables_from_xlsx_bytes,
    flatten_table_for_embedding,
    render_markdown_table,
)


# Keyword extraction/tokenizer tests have been moved to test_bm25_service.py.


# ============================================================================
# 2. Table extraction Tests
# ============================================================================

def test_markdown_generation_basic():
    md = render_markdown_table(
        ["Name", "Marks", "Grade"],
        [["Rohit", "95", "A"], ["Anil", "82", "B"]],
    )
    assert "| Name" in md
    assert "| -----" in md
    assert "| Rohit" in md


def test_flattened_generation_basic():
    flat = flatten_table_for_embedding(
        sheet="Students",
        headers=["Name", "Marks", "Grade"],
        rows=[["Rohit", "95", "A"], ["Anil", "82", "B"]],
    )
    assert "Sheet: Students" in flat
    assert "Row:" in flat
    assert "Marks = 95" in flat


def test_csv_extraction_headers_and_rows():
    data = b"Name,Marks,Grade\nRohit,95,A\nAnil,82,B\n"
    tables = extract_tables_from_csv_bytes(data, sheet_name="Sheet1", source_key="doc1")
    assert len(tables) == 1

    t = tables[0]
    assert t["sheet"] == "Sheet1"
    assert t["headers"] == ["Name", "Marks", "Grade"]
    assert t["rows"][0] == ["Rohit", "95", "A"]
    assert "| Name" in t["markdown"]
    assert "Marks = 95" in t["flattened_text"]


def test_xlsx_extraction_multiple_sheets():
    try:
        from openpyxl import Workbook
    except Exception as e:  # pragma: no cover
        pytest.skip(f"openpyxl not available: {e}")

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Students"
    ws1.append(["Name", "Marks", "Grade"])
    ws1.append(["Rohit", 95, "A"])
    ws1.append(["Anil", 82, "B"])

    ws2 = wb.create_sheet("Totals")
    ws2.append(["Metric", "Value"])
    ws2.append(["Average", 88.5])

    bio = io.BytesIO()
    wb.save(bio)

    tables = extract_tables_from_xlsx_bytes(bio.getvalue(), source_key="doc2")
    sheets = {t["sheet"] for t in tables}
    assert "Students" in sheets
    assert "Totals" in sheets


def test_docx_table_extraction_order():
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()

    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Name"
    t1.rows[0].cells[1].text = "Marks"
    t1.rows[1].cells[0].text = "Rohit"
    t1.rows[1].cells[1].text = "95"

    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Metric"
    t2.rows[0].cells[1].text = "Value"
    t2.rows[1].cells[0].text = "Total"
    t2.rows[1].cells[1].text = "177"

    bio = io.BytesIO()
    doc.save(bio)

    tables = extract_tables_from_docx_bytes(bio.getvalue(), source_key="doc3")
    assert len(tables) == 2

    assert tables[0]["headers"][:2] == ["Name", "Marks"]
    assert tables[1]["headers"][:2] == ["Metric", "Value"]


def test_docx_paragraph_extraction_normalizes_whitespace():
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    doc.add_paragraph("   Hello     world   ")
    doc.add_paragraph("   ")
    doc.add_paragraph("Line\t\twith\ninternal   whitespace")

    bio = io.BytesIO()
    doc.save(bio)

    text = extract_docx(bio.getvalue())
    assert "Hello world" in text
    assert "internal whitespace" in text
    assert "\n\n\n" not in text


def test_large_csv_extraction_completes_reasonably():
    import time

    rows = ["Name,Marks,Grade"]
    for i in range(10000):
        rows.append(f"Student{i},95,A")

    data = ("\n".join(rows)).encode("utf-8")

    start = time.perf_counter()

    tables = extract_tables_from_csv_bytes(
        data,
        sheet_name="Sheet1",
        source_key="perf_doc",
    )

    elapsed = time.perf_counter() - start

    assert len(tables) == 1
    assert len(tables[0]["rows"]) == 10000
    assert elapsed < 15


def test_corrupted_xlsx_returns_empty_tables():
    tables = extract_tables_from_xlsx_bytes(
        b"corrupted content",
        source_key="bad_doc",
    )
    assert tables == []


# ============================================================================
# 3. Hybrid indexing (text + tables)
# ============================================================================

class FakeIndexerCollection:
    def __init__(self):
        self.store = {}

    def upsert(self, embeddings, documents, metadatas, ids):
        for emb, doc, meta, _id in zip(embeddings, documents, metadatas, ids):
            self.store[_id] = {"embedding": emb, "document": doc, "metadata": meta}

    def get(self, where=None, ids=None):
        if where and "doc_id" in where:
            doc_id = where["doc_id"]
            result_ids, result_docs, result_metas = [], [], []
            for _id, item in self.store.items():
                if item["metadata"].get("doc_id") == doc_id:
                    result_ids.append(_id)
                    result_docs.append(item["document"])
                    result_metas.append(item["metadata"])
            return {"ids": result_ids, "documents": result_docs, "metadatas": result_metas}
        return {"ids": list(self.store.keys()), "documents": [], "metadatas": []}

    def delete(self, ids=None, **_kwargs):
        for _id in (ids or []):
            self.store.pop(_id, None)


@pytest.fixture
def fake_indexer_collection(monkeypatch):
    coll = FakeIndexerCollection()
    monkeypatch.setattr(indexer, "collection", coll)
    monkeypatch.setattr(pipeline, "collection", coll)
    return coll


@pytest.fixture(autouse=True)
def _no_node_push(monkeypatch):
    monkeypatch.setattr(indexer, "_push_chunks_to_node", lambda *_a, **_k: None)


@pytest.fixture(autouse=True)
def _mock_indexer_embeddings(monkeypatch):
    monkeypatch.setattr(pipeline, "embed_document", lambda *_a, **_k: [0.0, 0.1, 0.2])


def test_csv_indexes_table_chunks_with_metadata(fake_indexer_collection):
    data = b"Name,Marks,Grade\nRohit,95,A\nAnil,82,B\n"

    ok, added = indexer.index_bytes(
        doc_id="doc_csv",
        filename="students.csv",
        mimetype="text/csv",
        data=data,
        file_hash="h",
    )

    assert ok is True
    assert added >= 1

    metas = [v["metadata"] for v in fake_indexer_collection.store.values()]
    table_metas = [m for m in metas if m.get("is_table")]
    assert table_metas

    m0 = table_metas[0]
    assert m0.get("table_id")
    assert m0.get("table_index") == 0
    assert m0.get("doc_id") == "doc_csv"
    assert m0.get("row_start") == 0
    assert m0.get("row_end") == 2

    any_table_id = next(
        _id for _id, v in fake_indexer_collection.store.items() if v["metadata"].get("is_table")
    )
    stored_doc = fake_indexer_collection.store[any_table_id]["document"]
    stored_meta = fake_indexer_collection.store[any_table_id]["metadata"]
    assert "Row:" in stored_doc
    assert "|" not in stored_doc
    assert "markdown" in stored_meta
    assert "|" in stored_meta.get("markdown", "")


def test_table_deduplication_skips_duplicate_docx_tables(fake_indexer_collection):
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    for _ in range(2):
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "Name"
        t.rows[0].cells[1].text = "Marks"
        t.rows[1].cells[0].text = "Rohit"
        t.rows[1].cells[1].text = "95"

    bio = io.BytesIO()
    doc.save(bio)

    ok, added = indexer.index_bytes(
        doc_id="doc_docx_dupe",
        filename="dupe_tables.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=bio.getvalue(),
        file_hash="h3",
    )

    assert ok is True
    table_chunks = [v for v in fake_indexer_collection.store.values() if v["metadata"].get("is_table")]
    assert len(table_chunks) == 1


def test_chunk_id_uniqueness_uses_next_chunk_index(fake_indexer_collection, monkeypatch):
    monkeypatch.setattr(indexer, "split_sheet_sections", lambda _t: [(None, "body")])
    monkeypatch.setattr(chunking_mod, "chunk_text", lambda _b: ["A" * 400, "B" * 400])

    real_upsert = fake_indexer_collection.upsert

    def fake_flush_batch(collection_ref, batch_embeddings, batch_documents, batch_metadatas, batch_ids):
        real_upsert(batch_embeddings, batch_documents, batch_metadatas, batch_ids)
        return 0

    monkeypatch.setattr(pipeline, "_flush_batch", fake_flush_batch)
    monkeypatch.setattr(pipeline, "pack_blocks_into_chunks", lambda _sections: ["A" * 400, "B" * 400])

    monkeypatch.setattr(
        indexer,
        "extract_tables_for_file",
        lambda *_a, **_k: [
            {
                "table_id": "t1",
                "sheet": None,
                "headers": ["H1"],
                "rows": [["V1"]],
            }
        ],
    )

    ok, added = indexer.index_bytes(
        doc_id="doc_collision",
        filename="x.txt",
        mimetype="text/plain",
        data=b"body",
        file_hash="h4",
    )

    assert ok is True
    assert set(fake_indexer_collection.store.keys()) == {"doc_collision_0", "doc_collision_1", "doc_collision_2"}


def test_deterministic_chunk_numbering_consumes_indices_on_embedding_fail(fake_indexer_collection, monkeypatch):
    monkeypatch.setattr(indexer, "split_sheet_sections", lambda _t: [(None, "body")])
    monkeypatch.setattr(pipeline, "pack_blocks_into_chunks", lambda _sections: ["A" * 400, "B" * 400])

    calls = {"n": 0}

    def flaky_embeddings(*_a, **_k):
        calls["n"] += 1
        return None if calls["n"] == 1 else [0.0, 0.1, 0.2]

    monkeypatch.setattr(pipeline, "embed_document", flaky_embeddings)

    ok, added = indexer.index_text("doc_det", "sample.txt", "ignored")
    assert ok is True
    assert added == 1

    assert set(fake_indexer_collection.store.keys()) == {"doc_det_1"}
    meta = fake_indexer_collection.store["doc_det_1"]["metadata"]
    assert meta.get("chunk") == 1


def test_markdown_truncation_in_metadata(fake_indexer_collection, monkeypatch):
    monkeypatch.setattr(
        pipeline,
        "render_markdown_table",
        lambda *_a, **_k: ("| H |\n| --- |\n| V |\n" + ("x" * (indexer.MAX_MD_META_LEN + 200))),
        raising=False,
    )
    monkeypatch.setattr(
        pipeline,
        "flatten_table_for_embedding",
        lambda *_a, **_k: "Sheet: S1\n\nRow:\nH = V",
        raising=False,
    )

    headers = ["H"]
    rows = [["V"]]
    tables = [{"table_id": "t_big", "sheet": "S1", "headers": headers, "rows": rows}]

    chunk_records = []
    added = indexer._index_tables_spreadsheet(
        "doc_md",
        "big.xlsx",
        "xlsx",
        tables,
        start_chunk_index=0,
        chunk_records_out=chunk_records,
        file_hash="h_md",
    )

    assert added == 1
    stored = fake_indexer_collection.store["doc_md_0"]
    md_meta = stored["metadata"].get("markdown")
    assert isinstance(md_meta, str)
    assert len(md_meta) <= indexer.MAX_MD_META_LEN
    assert md_meta.endswith("...[truncated]")


def test_docx_table_only_is_indexed(fake_indexer_collection):
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Marks"
    t.rows[1].cells[0].text = "Rohit"
    t.rows[1].cells[1].text = "95"

    bio = io.BytesIO()
    doc.save(bio)

    ok, added = indexer.index_bytes(
        doc_id="doc_docx",
        filename="only_table.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=bio.getvalue(),
        file_hash="h2",
    )

    assert ok is True
    assert added >= 1

    metas = [v["metadata"] for v in fake_indexer_collection.store.values()]
    assert any(m.get("is_table") for m in metas)


# ============================================================================
# 4. Retrieval service query and ranking Tests
# ============================================================================

class FakeRetrievalCollection:
    def __init__(self, docs, dists, metas):
        self._docs = docs
        self._dists = dists
        self._metas = metas

    def query(self, *args, **kwargs):
        return {
            "documents": [self._docs],
            "distances": [self._dists],
            "metadatas": [self._metas],
        }


@pytest.fixture(autouse=True)
def _mock_retrieval_embeddings(monkeypatch):
    monkeypatch.setattr(retrieval_service, "embed_query", lambda _q: [0.0, 0.1, 0.2])


def test_table_question_boosts_table_chunks(monkeypatch):
    table_doc = "Sheet: Students\n\nRow:\nName = Rohit\nMarks = 95\nGrade = A"
    non_table_doc = ("This document explains grading policies and rules." * 2).strip()

    docs = [non_table_doc, table_doc]
    dists = [0.20, 0.25]
    metas = [
        {"is_table": False},
        {"is_table": True, "table_id": "t1", "table_index": 0},
    ]

    monkeypatch.setattr(
        retrieval_service,
        "collection",
        FakeRetrievalCollection(docs, dists, metas),
    )

    ctx, err = retrieval_service.retrieve_context(
        "What is the highest marks in the table?",
        "doc1",
    )

    assert err is None
    assert ctx is not None
    assert ctx.strip().startswith("Sheet: Students")


def test_generic_question_does_not_unfairly_boost_tables(monkeypatch):
    table_doc = "Sheet: Students\n\nRow:\nName = Rohit\nMarks = 95\nGrade = A"
    non_table_doc = ("This document explains grading policies and rules." * 2).strip()

    docs = [non_table_doc, table_doc]
    dists = [0.10, 0.20]
    metas = [
        {"is_table": False},
        {"is_table": True, "table_id": "t1", "table_index": 0},
    ]

    monkeypatch.setattr(
        retrieval_service,
        "collection",
        FakeRetrievalCollection(docs, dists, metas),
    )

    ctx, err = retrieval_service.retrieve_context(
        "Explain the grading policies and rules",
        "doc1",
    )

    assert err is None
    assert ctx is not None
    assert ctx.strip().startswith("This document explains grading policies")


# ============================================================================
# 5. Metadata Cache Tests (TTL, Cache Miss, Invalidation, Concurrency)
# ============================================================================

@pytest.fixture(autouse=True)
def _reset_meta_cache(monkeypatch):
    from services import retrieval_service
    with retrieval_service._doc_meta_cache_lock:
        retrieval_service._doc_meta_cache.clear()

    monkeypatch.setattr(
        retrieval_service,
        "_DOC_META_CACHE_TTL",
        300,
        raising=False,
    )
    yield
    with retrieval_service._doc_meta_cache_lock:
        retrieval_service._doc_meta_cache.clear()


def test_cache_miss_triggers_underlying_fetch(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}

    def fake_fetch(doc_id):
        calls["n"] += 1
        return {"contentHash": "h1"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    meta = retrieval_service.fetch_doc_meta_cached("doc1")

    assert meta.get("contentHash") == "h1"
    assert calls["n"] == 1


def test_cache_hit_avoids_repeated_fetch(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}

    def fake_fetch(doc_id):
        calls["n"] += 1
        return {"contentHash": "h1"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    meta1 = retrieval_service.fetch_doc_meta_cached("doc1")
    meta2 = retrieval_service.fetch_doc_meta_cached("doc1")

    assert meta1.get("contentHash") == "h1"
    assert meta2.get("contentHash") == "h1"
    assert calls["n"] == 1


def test_expired_entry_is_refreshed(monkeypatch):
    from services import retrieval_service

    now = {"t": 1000.0}

    def fake_time():
        return now["t"]

    monkeypatch.setattr(retrieval_service.time, "time", fake_time)

    calls = {"n": 0}

    def fake_fetch(doc_id):
        calls["n"] += 1
        return {"contentHash": f"h{calls['n']}"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    meta1 = retrieval_service.fetch_doc_meta_cached("doc1")
    assert meta1.get("contentHash") == "h1"
    assert calls["n"] == 1

    now["t"] += retrieval_service._DOC_META_CACHE_TTL + 1

    meta2 = retrieval_service.fetch_doc_meta_cached("doc1")
    assert meta2.get("contentHash") == "h2"
    assert calls["n"] == 2


def test_invalidate_cached_doc_meta_removes_entry(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}

    def fake_fetch(doc_id):
        calls["n"] += 1
        return {"contentHash": f"h{calls['n']}"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    assert retrieval_service.fetch_doc_meta_cached("doc1")["contentHash"] == "h1"

    retrieval_service.invalidate_cached_doc_meta("doc1")

    assert retrieval_service.fetch_doc_meta_cached("doc1")["contentHash"] == "h2"
    assert calls["n"] == 2


def test_blank_doc_id_is_handled_safely(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}

    def fake_fetch(_doc_id):
        calls["n"] += 1
        return {"contentHash": "h"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    assert retrieval_service.get_cached_doc_meta("") is None
    assert retrieval_service.fetch_doc_meta_cached("") == {"contentHash": "h"}
    assert calls["n"] == 1


def test_failed_fetch_is_cached_if_it_returns_empty_dict(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}

    def fake_fetch(doc_id):
        calls["n"] += 1
        return {}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    assert retrieval_service.fetch_doc_meta_cached("doc1") == {}
    assert retrieval_service.fetch_doc_meta_cached("doc1") == {}
    assert calls["n"] == 1


def test_concurrent_cache_miss_is_thread_safe(monkeypatch):
    from services import retrieval_service

    calls = {"n": 0}
    start_barrier = threading.Barrier(2)

    def fake_fetch(doc_id):
        start_barrier.wait()
        calls["n"] += 1
        return {"contentHash": "h1"}

    monkeypatch.setattr(
        retrieval_service,
        "fetch_doc_meta_from_node",
        fake_fetch,
    )

    results = []

    def worker():
        results.append(retrieval_service.fetch_doc_meta_cached("doc1"))

    t1 = threading.Thread(target=worker)
    t2 = threading.Thread(target=worker)

    t1.start()
    t2.start()
    t1.join(timeout=5)
    t2.join(timeout=5)

    assert not t1.is_alive()
    assert not t2.is_alive()

    assert len(results) == 2
    assert all(result.get("contentHash") == "h1" for result in results)

    assert 1 <= calls["n"] <= 2

    cached = retrieval_service.get_cached_doc_meta("doc1")
    assert cached is not None
    assert cached.get("contentHash") == "h1"


# ============================================================================
# 6. End-to-End PDF Indexing and Retrieval Integration Test
# ============================================================================

def test_end_to_end_pdf_indexing_and_retrieval(monkeypatch):
    # Create a shared fake collection simulating a real vector database
    fake_db = FakeIndexerCollection()
    
    def fake_query(query_embeddings, n_results, where=None, **kwargs):
        doc_id = where.get("doc_id") if where else None
        docs, dists, metas = [], [], []
        for item in fake_db.store.values():
            if not doc_id or item["metadata"].get("doc_id") == doc_id:
                docs.append(item["document"])
                dists.append(0.1)  # Simulated close distance
                metas.append(item["metadata"])
        return {
            "documents": [docs],
            "distances": [dists],
            "metadatas": [metas]
        }
        
    fake_db.query = fake_query
    
    monkeypatch.setattr(indexer, "collection", fake_db)
    monkeypatch.setattr(pipeline, "collection", fake_db)
    monkeypatch.setattr(retrieval_service, "collection", fake_db)
    
    # Mock PyMuPDF4LLM extraction response
    import utils.extraction as extraction_mod

    mock_pages = [
        {"page": 1, "text": "# Section A\nThis is page 1 content."},
        {"page": 2, "text": "# Section B\nThis is page 2 content explaining optimizer."}
    ]
    monkeypatch.setattr(extraction_mod, "extract_pdf", lambda _d: mock_pages)
    
    # Index the simulated PDF bytes
    ok, added = indexer.index_bytes(
        doc_id="end_to_end_doc",
        filename="manual.pdf",
        mimetype="application/pdf",
        data=b"pdf_bytes"
    )
    
    assert ok is True
    assert added >= 2
    
    # Retrieve context for a query targeting page 2 content
    ctx, err = retrieval_service.retrieve_context("optimizer", "end_to_end_doc")
    
    assert err is None
    assert ctx is not None
    # Ensure correct contextual retrieval and page contents return
    assert "page 2 content" in ctx


def test_retrieval_uses_embed_query():
    """Verify that retrieval path calls embed_query."""
    from unittest.mock import patch
    with patch("services.retrieval_service.embed_query") as mock_query:
        mock_query.return_value = [0.5, 0.5, 0.5]

        # Mock Chroma collection.query to avoid real DB access
        with patch("services.retrieval_service.collection") as mock_coll:
            mock_coll.query.return_value = {"documents": [[]], "distances": [[]], "metadatas": [[]]}

            # This should invoke embed_query
            retrieval_service.retrieve_context("User question", "doc123")

            mock_query.assert_called_once_with("User question")

